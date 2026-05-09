from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE = Path("C:/Programming/FYP")
TEMPLATE = BASE / "docs/reports/FYP-I-Report-Template-FA25-Research & Development.docx"
OUT = BASE / "docs/reports/Chapter_2_Literature_Review_Related_Work_Research_Gap.docx"


REFERENCES = [
    'Centers for Medicare & Medicaid Services, "Home Health Quality Reporting Program: Data Specifications," CMS, Page last modified Mar. 10, 2026. [Online]. Available: https://www.cms.gov/medicare/quality/home-health/data-specifications. Accessed: Apr. 29, 2026.',
    "Z. Wang et al., Measuring Documentation Burden in Healthcare, Technical Brief No. 47. Rockville, MD, USA: Agency for Healthcare Research and Quality, May 2024. [Online]. Available: https://www.ncbi.nlm.nih.gov/books/NBK608551/.",
    'U.S. Department of Health and Human Services, "The Security Rule," HHS.gov. [Online]. Available: https://www.hhs.gov/hipaa/for-professionals/security/index.html. Accessed: Apr. 29, 2026.',
    'R. Smith, "An overview of the Tesseract OCR engine," in Proc. 9th Int. Conf. Document Analysis and Recognition (ICDAR), 2007, pp. 629-633, doi: 10.1109/ICDAR.2007.4376991.',
    'A. R. Katti, C. Reisswig, C. Guder, S. Brarda, S. Bickel, J. Hoehne, and J. B. Faddoul, "Chargrid: Towards Understanding 2D Documents," in Proc. EMNLP, 2018, pp. 4459-4469. [Online]. Available: https://aclanthology.org/D18-1476/.',
    'Y. Xu, M. Li, L. Cui, S. Huang, F. Wei, and M. Zhou, "LayoutLM: Pre-training of Text and Layout for Document Image Understanding," in Proc. KDD, 2020, pp. 1192-1200, doi: 10.1145/3394486.3403172.',
    'Y. Huang, T. Lv, L. Cui, Y. Lu, and F. Wei, "LayoutLMv3: Pre-training for Document AI with Unified Text and Image Masking," arXiv:2204.08387, 2022. [Online]. Available: https://arxiv.org/abs/2204.08387.',
    'S. Appalaraju, B. Jasani, B. U. Kota, Y. Xie, and R. Manmatha, "DocFormer: End-to-End Transformer for Document Understanding," in Proc. ICCV, 2021; arXiv:2106.11539. [Online]. Available: https://arxiv.org/abs/2106.11539.',
    'G. Kim et al., "OCR-Free Document Understanding Transformer," in Proc. ECCV, 2022; arXiv:2111.15664. [Online]. Available: https://www.ecva.net/papers/eccv_2022/papers_ECCV/html/8042_ECCV_2022_paper.php.',
    'G. Bradski, "The OpenCV Library," Dr. Dobb\'s Journal of Software Tools, 2000.',
    'P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," in Advances in Neural Information Processing Systems, vol. 33, 2020. [Online]. Available: https://papers.nips.cc/paper/2020/hash/6b493230205f780e1bc26945df7481e5-Abstract.html.',
    'E. Asgari et al., "A framework to assess clinical safety and hallucination rates of LLMs for medical text summarisation," npj Digital Medicine, vol. 8, art. 274, 2025, doi: 10.1038/s41746-025-01670-7.',
    'K. H. Bowles and J. R. Cater, "Screening for risk of rehospitalization from home care: Use of the Outcomes Assessment Information Set and the Probability of Readmission instrument," Research in Nursing & Health, vol. 26, no. 2, pp. 118-127, Apr. 2003, doi: 10.1002/nur.10071.',
    'D. Kansagara et al., "Risk prediction models for hospital readmission: A systematic review," JAMA, vol. 306, no. 15, pp. 1688-1698, 2011, doi: 10.1001/jama.2011.1515.',
    'Y. Huang, A. Talwar, S. Chatterjee, and R. Aparasu, "Application of machine learning in predicting hospital readmissions: A scoping review of the literature," BMC Medical Research Methodology, vol. 21, art. 96, 2021, doi: 10.1186/s12874-021-01284-z.',
    'M. Hobensack, J. Song, D. Scharp, K. H. Bowles, and M. Topaz, "Machine learning applied to electronic health record data in home healthcare: A scoping review," International Journal of Medical Informatics, vol. 170, art. 104978, 2023, doi: 10.1016/j.ijmedinf.2022.104978.',
    'A. E. W. Johnson et al., "MIMIC-III, a freely accessible critical care database," Scientific Data, vol. 3, art. 160035, 2016, doi: 10.1038/sdata.2016.35.',
    'F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.',
    'T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. 22nd ACM SIGKDD Int. Conf. Knowledge Discovery and Data Mining, 2016, pp. 785-794, doi: 10.1145/2939672.2939785.',
    'G. Ke et al., "LightGBM: A Highly Efficient Gradient Boosting Decision Tree," in Advances in Neural Information Processing Systems, vol. 30, 2017, pp. 3146-3154. [Online]. Available: https://papers.nips.cc/paper/6907-lightgbm-a-highly-efficient-gradient-boosting-decision-tree.',
    'S. M. Lundberg and S.-I. Lee, "A Unified Approach to Interpreting Model Predictions," in Advances in Neural Information Processing Systems, vol. 30, 2017. [Online]. Available: https://papers.nips.cc/paper/7062-a.',
    'National Institute of Standards and Technology, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," NIST, Jan. 2023. [Online]. Available: https://www.nist.gov/itl/ai-risk-management-framework. Accessed: Apr. 29, 2026.',
]


LIT_SECTIONS = [
    (
        "Clinical Documentation and Home Health Context",
        [
            "Home health and hospice agencies depend on accurate documentation because assessment records affect patient evaluation, care coordination, quality reporting, and reimbursement. CMS identifies the Outcome and Assessment Information Set (OASIS) as the data collection instrument used by home health agencies, and the current OASIS-E2 data submission specifications are scheduled for implementation on April 1, 2026 [1]. For the proposed project, this means the system cannot treat OASIS and Plan of Care (POC) forms as ordinary text documents; it must preserve field meaning, assessment timing, clinical context, and review traceability.",
            "The proposal identifies a practical operational problem: clinicians and administrators repeatedly read scanned or semi-structured forms, re-enter patient information, validate codes and dates, and prepare care-plan content manually. This aligns with broader evidence that clinical documentation burden is a measurable health informatics problem. AHRQ's 2024 technical brief reports that documentation burden is commonly measured through EHR time, clinical documentation activity, after-hours work, administrative tasks, workflow fragmentation, and usability measures, with limited evidence that current measures generalize cleanly across settings [2]. Therefore, the proposed FYP should evaluate not only model accuracy but also reviewer time, correction workload, and workflow usability.",
            "Because the project handles protected health information, security cannot be added after model development. The HIPAA Security Rule requires administrative, physical, and technical safeguards for electronic protected health information [3]. In system terms, this supports role-based access control, encryption at rest and in transit, secure secret management, audit logging, and strict control over where patient documents and model outputs are stored.",
        ],
    ),
    (
        "OCR and Key Information Extraction",
        [
            "Traditional OCR converts document images into machine-readable text. Tesseract remains a widely cited open-source OCR baseline and is useful for printed text experiments and cost-controlled local processing [4]. For a student FYP, this is the most practical primary OCR engine because it can be installed locally, tested without cloud fees, and integrated with Python. However, OCR alone does not solve the FYP problem because OASIS and POC documents contain field labels, checkboxes, tables, signatures, dates, diagnosis codes, and spatial relationships. A raw text stream may lose the association between a label and its value, so post-processing is needed to produce structured JSON outputs.",
            "Image preprocessing is therefore important before OCR. OpenCV provides practical computer-vision operations such as grayscale conversion, denoising, thresholding, deskewing, contour detection, and cropping [10]. In this project, OpenCV can improve low-quality scanned OASIS-E2 and POC forms before Tesseract reads them. The output can then be parsed with field dictionaries, regular expressions, and validation rules.",
            "Early layout-aware approaches address this limitation by preserving document geometry. Chargrid represents each document page as a two-dimensional character grid and uses visual document structure for information extraction [5]. This is important for scanned healthcare forms because fields are often defined by location and neighboring labels. However, such methods still require domain adaptation and annotated data to perform reliably on clinical forms with variable scan quality, handwriting, and local template changes.",
        ],
    ),
    (
        "Layout-Aware and Multimodal Document AI",
        [
            "Transformer-based document models extend OCR pipelines by learning relationships among text tokens, bounding boxes, and visual layout. LayoutLM jointly models text and layout for scanned document understanding and demonstrated that layout information improves information extraction from forms [6]. LayoutLMv3 further unifies text and image masking, making document pre-training more general across text-centric and image-centric document tasks [7]. DocFormer similarly combines text, vision, and spatial features through multimodal self-attention for visual document understanding [8].",
            "OCR-free models such as Donut move in a different direction by mapping document images directly to structured outputs and reducing dependence on external OCR systems [9]. This is relevant to low-quality scans and multilingual environments, but it also raises deployment questions: domain-specific fine-tuning, ground-truth annotation, compute cost, and quality assurance are still required. For a final year project MVP, these models should be treated as future research extensions rather than the primary implementation path.",
            "Commercial document AI services can be useful benchmarks, but they are not the preferred student implementation route because cloud usage may create cost, account, and data-handling constraints. The practical MVP should therefore use Tesseract OCR, OpenCV preprocessing, rule-based extraction, and schema validation first, while leaving managed cloud document AI as an optional future extension.",
        ],
    ),
    (
        "RAG-Based Plan of Care Drafting",
        [
            "The proposed POC module should not generate unsupported clinical statements. Retrieval-Augmented Generation (RAG) is therefore a core implementation component, not a future enhancement. RAG conditions generation on retrieved evidence [11], so the MVP will index reviewer-approved OASIS-E2 fields, extracted POC snippets, diagnosis codes, medication notes, care-plan templates, and source snippets. The generator will then draft POC sections only from retrieved evidence and approved structured fields.",
            "For a student-friendly implementation, the retrieval layer can start with a local vector index such as FAISS or Chroma, or a TF-IDF/BM25 fallback if embeddings are difficult to run. Embeddings can use open-source sentence-transformer models where hardware allows. The generation layer can use a constrained local model, an approved API model, or a structured prompt/template wrapper, but the project requirement is that every generated section must include citations to retrieved snippets, missing-evidence warnings, and mandatory clinician approval.",
            "Recent clinical summarization research shows why uncontrolled LLM generation is unsafe. LLM-generated clinical text must be evaluated for hallucinations and omissions because unsupported or missing information can affect patient safety [12]. Therefore, this project treats RAG as a controlled drafting assistant: retrieval, citations, validation rules, unsupported-statement checks, and human review are all part of the core MVP.",
        ],
    ),
    (
        "Readmission Risk Prediction and Explainability",
        [
            "Hospital readmission risk prediction is a mature but difficult area. Bowles and Cater studied rehospitalization risk in home care using OASIS and the Probability of Readmission instrument, showing that OASIS-derived function scores can help identify at-risk home care patients but may not be sufficient alone [13]. A later systematic review of readmission prediction models found that many models had only moderate discrimination and that preventable readmission was rarely addressed directly [14]. These findings support the proposal's decision to start with a baseline, calibrated 30-day readmission model rather than presenting risk prediction as a solved problem.",
            "Machine-learning reviews show that EHR, claims, and population data are frequently used for readmission prediction, with tree-based methods, neural networks, regularized logistic regression, and support vector machines appearing commonly in the literature [15]. In home healthcare specifically, Hobensack et al. found that many studies used standardized assessments such as OASIS, hospitalization was the most common outcome, tree-based algorithms were frequent, and many studies had high or unclear risk of bias [16]. This indicates that the FYP should report AUROC together with calibration metrics such as Brier score, use a held-out test split, and document data limitations.",
            "Public clinical datasets such as MIMIC-III are valuable for experimentation because they provide de-identified EHR data, notes, medications, procedures, and outcomes for a large critical-care population [17]. However, MIMIC-III is not a home health dataset, so any model trained or tested on it must be described as a baseline or methodological prototype, not as a validated home health deployment model. For student implementation, scikit-learn Logistic Regression and Random Forest should be used first because they are free, well documented, and easy to evaluate [18]. XGBoost and LightGBM can be added as optional stronger tabular baselines if time allows [19], [20]. SHAP can then provide local feature-attribution explanations so reviewers can see which features contributed to a risk score [21].",
        ],
    ),
    (
        "Human-in-the-Loop Governance",
        [
            "The literature supports a human-in-the-loop design for this project. OCR can pre-fill fields, RAG can draft source-cited care-plan sections, and machine-learning models can estimate readmission risk, but clinicians must remain responsible for final approval. NIST's AI Risk Management Framework emphasizes trustworthiness considerations during AI design, development, use, and evaluation [22]. For the proposed system, this means model outputs should be visible, editable, logged, and measurable. The MVP should therefore focus on a safe clinician workflow: upload, preprocess, extract, review, approve, generate, and predict.",
        ],
    ),
]


RELATED_WORK_ROWS = [
    ("[1]", "CMS OASIS-E2 data specifications", "Home health assessment and quality reporting", "Defines the standardized OASIS data submission structure and implementation timeline.", "Regulatory/assessment specification, not an automation method.", "Defines the form family and structured fields that the MVP must support."),
    ("[2]", "AHRQ documentation burden technical brief", "Clinical documentation workload", "Synthesizes measurement approaches for documentation burden across healthcare settings.", "Evidence does not fully generalize across all roles and settings.", "Supports reviewer-time and workflow-burden KPIs, not only extraction accuracy."),
    ("[4]", "Tesseract OCR", "Baseline text recognition", "Shows that open-source OCR can convert printed document images into text.", "Does not preserve enough clinical field semantics or layout relationships by itself.", "Useful as a low-cost OCR baseline or fallback for demo/local processing."),
    ("[5]", "Chargrid", "2D document representation", "Preserves document layout through a character-grid representation for structured document understanding.", "Requires task-specific training and is not healthcare-form specific.", "Motivates use of layout information for OASIS/POC field extraction."),
    ("[6]", "LayoutLM", "Text-layout pre-training", "Combines text tokens and bounding-box layout information for scanned-document understanding.", "Depends on OCR quality and downstream fine-tuning data.", "Supports layout-aware KIE for forms, checkboxes, and tables."),
    ("[7]", "LayoutLMv3", "Unified text and image masking", "Uses unified multimodal pre-training for broader document AI tasks.", "Still needs adaptation to the exact clinical schema and scan conditions.", "Candidate future model after the MVP extraction schema is stable."),
    ("[8]", "DocFormer", "Multimodal visual document understanding", "Combines text, vision, and spatial features with multimodal attention.", "Higher implementation complexity for an FYP MVP.", "Useful benchmark for comparing document-understanding design choices."),
    ("[9]", "Donut", "OCR-free document understanding", "Generates structured outputs directly from document images without external OCR.", "Needs domain data and careful validation to avoid malformed or unsupported outputs.", "Possible extension for poor OCR cases after baseline pipeline works."),
    ("[10]", "OpenCV preprocessing", "Computer vision for scanned documents", "Supports grayscale conversion, thresholding, denoising, deskewing, contour detection, and image cleanup before OCR.", "Preprocessing improves OCR but does not understand clinical meaning by itself.", "Core student-friendly component before Tesseract OCR."),
    ("[11], [12]", "RAG and clinical LLM safety work", "Core grounded generation and hallucination control", "Retrieval can ground generation in source evidence; clinical note studies show hallucinations and omissions require explicit evaluation.", "RAG still needs retrieval evaluation, citation checks, unsupported-statement detection, and clinician approval.", "Defines the core POC drafting approach: retrieve approved evidence, generate a cited draft, and route it through human review."),
    ("[13]", "Bowles and Cater home-care rehospitalization study", "OASIS and readmission screening", "Compares OASIS-related scores with the Probability of Readmission instrument for home-care rehospitalization prediction.", "Older study and not designed around modern ML or current OASIS versions.", "Shows that OASIS-derived features can contribute to home-care risk identification."),
    ("[14], [15]", "Readmission prediction reviews", "General hospital readmission modeling", "Summarize model families, data sources, and common limitations in readmission prediction.", "Many models show moderate performance and limited transferability.", "Justifies using AUROC plus calibration and local validation rather than accuracy alone."),
    ("[16]", "Hobensack et al. HHC ML scoping review", "Machine learning in home healthcare", "Finds frequent use of standardized assessments, hospitalization outcomes, and tree-based algorithms in HHC ML studies.", "Many studies have high or unclear risk of bias and incomplete psychosocial/interpersonal predictors.", "Directly informs feature engineering and risk documentation for the FYP."),
    ("[18]-[21]", "scikit-learn, XGBoost, LightGBM, and SHAP", "Student-friendly tabular risk modeling", "Logistic Regression and Random Forest provide implementable baselines; gradient boosting can improve performance; SHAP explains individual predictions.", "Explainability does not fix data bias, calibration error, or workflow misuse.", "Supports a transparent readmission-risk module with clinician-facing explanations."),
    ("[3], [22]", "HIPAA Security Rule and NIST AI RMF", "Security and AI governance", "Establishes safeguards for ePHI and risk-management expectations for trustworthy AI systems.", "Frameworks do not implement controls automatically.", "Guides RBAC, encryption, audit logs, model-output traceability, and human approval."),
]


GAPS = [
    ("Gap 1: Generic document AI is not enough for OASIS and POC workflows.", "Existing OCR, KIE, and document transformer research performs well on general forms, receipts, and benchmark document tasks, but it does not directly solve OASIS/POC field semantics, healthcare-specific validation rules, clinical coding consistency, or Plan of Care traceability. The FYP addresses this by defining a project-specific extraction schema, field dictionary, validation rules, and reviewer workflow for home health and hospice documentation."),
    ("Gap 2: Prior work often optimizes model accuracy without closing the review loop.", "Many document-understanding systems stop at extraction results. In healthcare documentation, an extracted value is not final until a clinician reviews and approves it. The proposed system fills this gap by keeping a human-in-the-loop interface with accept, edit, reject, confidence display, source snippets, and audit logs."),
    ("Gap 3: Care-plan generation requires controlled RAG, not free-form generation.", "RAG improves factual grounding, but clinical LLM outputs can still contain hallucinations or omissions if retrieval and citations are weak. A safe student MVP should implement RAG with approved OASIS-E2 fields, source snippets, local retrieval, section-level citations, missing-evidence warnings, and clinician approval. The FYP contributes a source-grounded POC draft workflow rather than an autonomous care-plan generator."),
    ("Gap 4: Readmission models are not automatically transferable to home healthcare.", "General readmission studies frequently use hospital EHR, claims, or critical-care data. Home health populations have different documentation patterns, assessment timings, functional measures, caregiver context, and post-discharge risks. The FYP should therefore treat MIMIC-style datasets and generic readmission studies as methodological baselines while building a home-health feature set around OASIS fields, ADL/mobility indicators, diagnosis codes, prior utilization, medication burden, and clinician-reviewed data."),
    ("Gap 5: Calibration and explainability are under-addressed in student prototypes.", "A high AUROC is insufficient for clinical workflow. A risk score must be calibrated, explainable, and monitored for error patterns. The proposed project addresses this by targeting AUROC, Brier score, confidence intervals, error distributions, and SHAP-style explanations for individual scores."),
    ("Gap 6: Security and auditability must be part of the system architecture.", "Healthcare AI prototypes often demonstrate extraction or prediction but leave PHI controls, role separation, secret handling, and audit logs vague. This project addresses that gap by making RBAC, encryption, private storage, audit logging, and no-hardcoded-secret rules part of the MVP definition rather than final-stage polish."),
    ("Gap 7: End-to-end evaluation is missing from isolated research components.", "The related work contains strong individual components: OCR, image preprocessing, layout-aware extraction, RAG-based generation, risk models, and explainability. The research gap is the absence of an integrated, measurable workflow for home health documentation. The FYP fills this by evaluating the complete chain: upload -> preprocessing -> OCR -> extraction -> review -> approval -> retrieval -> cited POC generation -> readmission score, using KPIs such as field-level F1, reviewer time reduction, citation coverage, section-level POC approval, risk calibration, latency, and audit completeness."),
]


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def style_paragraph(paragraph, size=11, before=0, after=6, line=1.0, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_para(doc, text="", style=None, size=11, before=0, after=6, align=None):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size)
    style_paragraph(paragraph, size=size, before=before, after=after, align=align)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    size = 16 if level == 1 else 14 if level == 2 else 12
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=True, italic=(level == 3), color=(0, 0, 0))
    paragraph.paragraph_format.space_before = Pt(12 if level > 1 else 0)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def set_cell_text(cell, text, bold=False, size=8, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(width.inches * 1440) for width in widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(col)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def load_table_font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for paragraph in str(text).split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_related_work_table(rows, output_path):
    widths = [80, 260, 370, 520]
    margin = 18
    padding = 12
    header_font = load_table_font(22, bold=True)
    body_font = load_table_font(20)
    line_gap = 5
    headers = ["Ref.", "Related work", "Method / contribution", "Limitation and FYP relevance"]

    probe = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(probe)

    rendered_rows = []
    all_rows = [headers] + [
        [
            row[0],
            f"{row[1]}\nFocus: {row[2]}",
            row[3],
            f"Limitation: {row[4]}\nRelevance: {row[5]}",
        ]
        for row in rows
    ]
    for row_index, row in enumerate(all_rows):
        font = header_font if row_index == 0 else body_font
        wrapped_cells = [wrap_text(draw, cell, font, widths[i] - 2 * padding) for i, cell in enumerate(row)]
        line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
        height = max(len(lines) * line_height for lines in wrapped_cells) + 2 * padding
        rendered_rows.append((wrapped_cells, font, line_height, max(height, 58 if row_index == 0 else 70)))

    img_width = sum(widths) + 2 * margin
    img_height = sum(row[3] for row in rendered_rows) + 2 * margin
    image = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(image)

    y = margin
    for row_index, (wrapped_cells, font, line_height, row_height) in enumerate(rendered_rows):
        x = margin
        fill = "#D9EAF7" if row_index == 0 else "#FFFFFF"
        draw.rectangle([margin, y, margin + sum(widths), y + row_height], fill=fill)
        for i, lines in enumerate(wrapped_cells):
            x2 = x + widths[i]
            draw.rectangle([x, y, x2, y + row_height], outline="#4F5B66", width=2 if row_index == 0 else 1)
            text_y = y + padding
            for line in lines:
                draw.text((x + padding, text_y), line, fill="black", font=font)
                text_y += line_height
            x = x2
        y += row_height

    image.save(output_path, "PNG")


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def clear_headers_and_footers(doc):
    for section in doc.sections:
        for part in [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]:
            part.is_linked_to_previous = False
            for child in list(part._element):
                part._element.remove(child)


def main():
    doc = Document(TEMPLATE)
    clear_body(doc)
    set_portrait(doc.sections[0])

    style_names = {style.name for style in doc.styles}
    for style_name in ["Normal", "Normal (Web)", "Body Text"]:
        if style_name in style_names:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style.font.size = Pt(11)

    title = add_para(doc, "Chapter 2 Draft Sections", size=16, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    add_para(
        doc,
        "Literature Review, Summary Table of Related Work, and Research Gap Identification",
        size=12,
        after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    meta = add_para(
        doc,
        "Project base: HIPAA-Compliant Home Health and Hospice AI Platform. Source context: existing FYP proposal, FYP-I R&D report template, and report guidelines. Prepared: April 29, 2026.",
        size=10,
        after=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for run in meta.runs:
        run.italic = True

    add_heading(doc, "2.1 Literature Review", 2)
    add_para(
        doc,
        "This section reviews literature and technical work relevant to the proposed platform: secure document upload, OCR and key information extraction from OASIS/POC forms, human review, source-cited Plan of Care drafting, 30-day readmission risk prediction, and audit-ready security controls.",
    )
    for title_text, paragraphs in LIT_SECTIONS:
        add_heading(doc, title_text, 3)
        for text in paragraphs:
            add_para(doc, text)

    doc.add_page_break()
    add_heading(doc, "2.2 Summary Table of Related Work", 2)
    add_para(
        doc,
        "Table 2.1 summarizes the most relevant prior work and explains how each source informs the proposed FYP system.",
        size=10,
    )
    caption_style = "Caption" if "Caption" in style_names else None
    caption = add_para(
        doc,
        "Table 2.1. Summary of related work for AI-assisted home health and hospice documentation.",
        style=caption_style,
        size=10,
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for run in caption.runs:
        run.bold = True

    with TemporaryDirectory() as temp_dir:
        chunks = [RELATED_WORK_ROWS[0:5], RELATED_WORK_ROWS[5:10], RELATED_WORK_ROWS[10:15]]
        for idx, chunk in enumerate(chunks, start=1):
            image_path = Path(temp_dir) / f"table_2_1_part_{idx}.png"
            draw_related_work_table(chunk, image_path)
            image_paragraph = doc.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_run = image_paragraph.add_run()
            image_run.add_picture(str(image_path), width=Inches(6.45))
            if idx < len(chunks):
                cont = add_para(
                    doc,
                    f"Table 2.1 continued on next page (part {idx + 1}).",
                    size=9,
                    after=2,
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                )
                cont.runs[0].italic = True
                doc.add_page_break()

    note = add_para(
        doc,
        "The table shows that individual research streams are mature, but the FYP contribution is their controlled integration into a secure, reviewable workflow for home health and hospice documentation.",
        size=10,
        before=8,
        after=0,
    )
    note.runs[0].italic = True

    doc.add_page_break()
    add_heading(doc, "2.3 Research Gap Identification", 2)
    add_para(
        doc,
        "The reviewed work supports the feasibility of each technical component, but it also exposes gaps that justify the proposed FYP. The main gap is not the absence of OCR, image preprocessing, rule-based extraction, RAG-based generation, or readmission models individually; it is the absence of a secure, domain-specific, human-reviewed, end-to-end workflow for OASIS and Plan of Care documentation in home health and hospice settings.",
    )
    for gap_title, body in GAPS:
        paragraph = add_para(doc, gap_title, before=4, after=2)
        paragraph.runs[0].bold = True
        add_para(doc, body)
    add_para(
        doc,
        "Overall research gap statement: Existing work provides strong building blocks for OCR, image preprocessing, layout-aware extraction, RAG-based generation, risk modeling, explainability, and AI governance. However, current literature does not provide a complete, HIPAA-aware and student-implementable workflow that converts OASIS/POC document images into reviewed structured data, source-cited RAG care-plan drafts, and calibrated readmission-risk outputs while preserving auditability. This FYP addresses that gap by designing and evaluating an integrated clinician-centered platform with measurable extraction, retrieval, citation, review, generation, prediction, and security outcomes.",
        before=6,
        after=10,
    )

    add_heading(doc, "References Used in This Draft", 2)
    add_para(
        doc,
        "These IEEE-style references should be merged into the final report reference list in first-citation order when these sections are inserted into the full report.",
        size=10,
    )
    for idx, ref in enumerate(REFERENCES, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.3)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(f"[{idx}] {ref}")
        set_run_font(run, size=9)

    clear_headers_and_footers(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
